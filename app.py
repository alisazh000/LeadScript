import os
import io
import json
import sqlite3
import secrets
import hashlib
import zipfile
import datetime as dt
from functools import wraps
from collections import defaultdict, deque

from flask import Flask, jsonify, request, g, send_file, render_template
from werkzeug.security import generate_password_hash, check_password_hash
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.getenv("LEADSCRIPT_DB_PATH", os.path.join(BASE_DIR, "leadscript.db"))


def create_app() -> Flask:
    app = Flask(__name__, static_folder="static", template_folder="templates")
    app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", secrets.token_hex(32))
    app.config["TOKEN_MAX_AGE"] = int(os.getenv("TOKEN_MAX_AGE", "604800"))
    app.config["RATE_LIMIT_PER_MINUTE"] = int(os.getenv("RATE_LIMIT_PER_MINUTE", "120"))

    init_db()
    limiter = SimpleRateLimiter(app.config["RATE_LIMIT_PER_MINUTE"])

    @app.after_request
    def set_security_headers(resp):
        resp.headers["X-Content-Type-Options"] = "nosniff"
        resp.headers["X-Frame-Options"] = "DENY"
        resp.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        resp.headers["Permissions-Policy"] = "geolocation=(), microphone=(), camera=()"
        resp.headers["Cross-Origin-Opener-Policy"] = "same-origin"
        resp.headers["Cross-Origin-Resource-Policy"] = "same-origin"
        # CSP tuned for same-origin assets only
        resp.headers["Content-Security-Policy"] = (
            "default-src 'self'; "
            "script-src 'self'; "
            "style-src 'self'; "
            "img-src 'self' data:; "
            "connect-src 'self'; "
            "font-src 'self' data:; "
            "object-src 'none'; "
            "base-uri 'self'; "
            "frame-ancestors 'none'"
        )
        return resp

    @app.before_request
    def apply_rate_limit():
        # No throttling for static assets.
        if request.path.startswith("/static/"):
            return
        ip = request.headers.get("X-Forwarded-For", request.remote_addr or "unknown").split(",")[0].strip()
        if not limiter.allow(ip):
            return jsonify({"error": "Too many requests"}), 429

    @app.route("/")
    def index():
        return render_template("index.html")

    @app.get("/download/leadscript.zip")
    def download_app_zip():
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            include_files = [
                "app.py",
                "requirements.txt",
                "README.md",
                ".env.example",
                "render.yaml",
                "Dockerfile",
            ]
            for rel in include_files:
                abs_path = os.path.join(BASE_DIR, rel)
                if os.path.isfile(abs_path):
                    zf.write(abs_path, arcname=rel)

            for folder in ("templates", "static", "tests"):
                folder_path = os.path.join(BASE_DIR, folder)
                if not os.path.isdir(folder_path):
                    continue
                for root, _dirs, files in os.walk(folder_path):
                    for name in files:
                        p = os.path.join(root, name)
                        arc = os.path.relpath(p, BASE_DIR)
                        zf.write(p, arcname=arc)

        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name="LeadScript.zip",
            mimetype="application/zip",
        )

    @app.get("/download/LeadScript.exe")
    def download_windows_exe():
        exe_path = os.path.join(BASE_DIR, "dist", "LeadScript.exe")
        if not os.path.isfile(exe_path):
            return jsonify({"error": "Windows build not found"}), 404
        return send_file(
            exe_path,
            as_attachment=True,
            download_name="LeadScript.exe",
            mimetype="application/vnd.microsoft.portable-executable",
        )

    @app.route("/privacy")
    def privacy():
        return render_template("privacy.html")

    @app.route("/terms")
    def terms():
        return render_template("terms.html")

    @app.get("/api/health")
    def health():
        return jsonify({"ok": True, "time": dt.datetime.utcnow().isoformat() + "Z"})

    @app.post("/api/auth/register")
    def register():
        body = request.get_json(silent=True) or {}
        email = (body.get("email") or "").strip().lower()
        password = body.get("password") or ""
        name = (body.get("name") or "").strip() or "Creator"
        if not email or "@" not in email:
            return jsonify({"error": "Invalid email"}), 400
        if len(password) < 8:
            return jsonify({"error": "Password must be at least 8 chars"}), 400

        db = get_db()
        cur = db.execute("SELECT id FROM users WHERE email = ?", (email,))
        if cur.fetchone():
            return jsonify({"error": "Email already exists"}), 409

        pwd_hash = generate_password_hash(password)
        is_admin = 1 if is_admin_email(email) else 0
        db.execute(
            "INSERT INTO users(email,name,password_hash,is_admin,created_at) VALUES (?,?,?,?,?)",
            (email, name, pwd_hash, is_admin, utcnow()),
        )
        db.commit()
        user = db.execute("SELECT id,email,name,is_admin FROM users WHERE email = ?", (email,)).fetchone()
        token = sign_token(
            app,
            {
                "uid": user["id"],
                "email": user["email"],
                "name": user["name"],
                "is_admin": int(user["is_admin"] or 0),
            },
        )
        return jsonify({"token": token, "user": dict(user)})

    @app.post("/api/auth/login")
    def login():
        body = request.get_json(silent=True) or {}
        email = (body.get("email") or "").strip().lower()
        password = body.get("password") or ""
        db = get_db()
        user = db.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
        if not user or not check_password_hash(user["password_hash"], password):
            return jsonify({"error": "Invalid credentials"}), 401
        token = sign_token(
            app,
            {
                "uid": user["id"],
                "email": user["email"],
                "name": user["name"],
                "is_admin": int(user["is_admin"] or 0),
            },
        )
        return jsonify(
            {
                "token": token,
                "user": {
                    "id": user["id"],
                    "email": user["email"],
                    "name": user["name"],
                    "is_admin": int(user["is_admin"] or 0),
                },
            }
        )

    @app.post("/api/auth/guest")
    def guest_login():
        db = get_db()
        guest_tag = secrets.token_hex(4)
        email = f"guest_{guest_tag}@leadscript.local"
        name = f"Guest-{guest_tag}"
        # Random internal password hash for a non-interactive guest account.
        pwd_hash = generate_password_hash(secrets.token_urlsafe(24))
        db.execute(
            "INSERT INTO users(email,name,password_hash,is_admin,created_at) VALUES (?,?,?,?,?)",
            (email, name, pwd_hash, 0, utcnow()),
        )
        db.commit()
        user = db.execute("SELECT id,email,name,is_admin FROM users WHERE email = ?", (email,)).fetchone()
        token = sign_token(
            app,
            {
                "uid": user["id"],
                "email": user["email"],
                "name": user["name"],
                "is_admin": 0,
            },
        )
        return jsonify({"token": token, "user": dict(user), "guest": True})

    @app.get("/api/auth/me")
    @auth_required(app)
    def me():
        return jsonify({"user": g.user})

    @app.get("/api/projects")
    @auth_required(app)
    def list_projects():
        db = get_db()
        rows = db.execute(
            "SELECT id,title,status,tags,created_at,updated_at FROM projects WHERE user_id = ? ORDER BY updated_at DESC",
            (g.user["id"],),
        ).fetchall()
        return jsonify([row_to_project(r) for r in rows])

    @app.post("/api/projects")
    @auth_required(app)
    def create_project():
        body = request.get_json(silent=True) or {}
        title = (body.get("title") or "Untitled Project").strip()
        status = (body.get("status") or "draft").strip()
        tags = body.get("tags") or []
        tags_json = json.dumps(tags, ensure_ascii=False)
        now = utcnow()
        db = get_db()
        db.execute(
            "INSERT INTO projects(user_id,title,status,tags,created_at,updated_at) VALUES (?,?,?,?,?,?)",
            (g.user["id"], title, status, tags_json, now, now),
        )
        db.commit()
        row = db.execute(
            "SELECT id,title,status,tags,created_at,updated_at FROM projects WHERE rowid = last_insert_rowid()"
        ).fetchone()
        return jsonify(row_to_project(row)), 201

    @app.delete("/api/projects/<int:project_id>")
    @auth_required(app)
    def delete_project(project_id: int):
        db = get_db()
        owns = db.execute("SELECT id FROM projects WHERE id = ? AND user_id = ?", (project_id, g.user["id"])).fetchone()
        if not owns:
            return jsonify({"error": "Project not found"}), 404
        db.execute("DELETE FROM chat_messages WHERE project_id = ?", (project_id,))
        db.execute("DELETE FROM scripts WHERE project_id = ?", (project_id,))
        db.execute("DELETE FROM projects WHERE id = ?", (project_id,))
        db.commit()
        return jsonify({"ok": True})

    @app.get("/api/projects/<int:project_id>/chats")
    @auth_required(app)
    def list_chats(project_id: int):
        if not ensure_project_owner(project_id, g.user["id"]):
            return jsonify({"error": "Project not found"}), 404
        rows = get_db().execute(
            "SELECT id,role,content,created_at FROM chat_messages WHERE project_id = ? ORDER BY id ASC",
            (project_id,),
        ).fetchall()
        return jsonify([dict(r) for r in rows])

    @app.post("/api/chats/message")
    @auth_required(app)
    def chat_message():
        body = request.get_json(silent=True) or {}
        project_id = int(body.get("project_id") or 0)
        text = (body.get("text") or "").strip()
        style = body.get("style") or "balanced"
        detail = body.get("detail") or "standard"
        focus = body.get("focus") or "content"
        creativity = body.get("creativity") or "medium"
        model = (body.get("model") or os.getenv("OPENAI_MODEL", "gpt-5.3-codex")).strip()
        temperature = clamp_float(body.get("temperature"), 0.0, 2.0, 0.7)
        max_tokens = clamp_int(body.get("max_tokens"), 80, 1500, 350)
        response_mode = (body.get("response_mode") or "normal").strip()
        system_prompt = (body.get("system_prompt") or "").strip()
        memory_enabled = bool(body.get("memory_enabled", True))
        memory_window = clamp_int(body.get("memory_window"), 2, 30, 8)
        if not text:
            return jsonify({"error": "Empty message"}), 400
        if not ensure_project_owner(project_id, g.user["id"]):
            return jsonify({"error": "Project not found"}), 404

        db = get_db()
        history = []
        if memory_enabled:
            rows = db.execute(
                "SELECT role,content FROM chat_messages WHERE project_id = ? ORDER BY id DESC LIMIT ?",
                (project_id, memory_window * 2),
            ).fetchall()
            history = [dict(r) for r in reversed(rows)]

        ai_reply = generate_ai_reply(
            text=text,
            style=style,
            detail=detail,
            focus=focus,
            creativity=creativity,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            response_mode=response_mode,
            system_prompt=system_prompt,
            history=history,
        )
        now = utcnow()
        db.execute(
            "INSERT INTO chat_messages(project_id,role,content,created_at) VALUES (?,?,?,?)",
            (project_id, "user", text, now),
        )
        db.execute(
            "INSERT INTO chat_messages(project_id,role,content,created_at) VALUES (?,?,?,?)",
            (project_id, "assistant", ai_reply, now),
        )
        db.execute("UPDATE projects SET updated_at = ? WHERE id = ?", (now, project_id))
        db.commit()
        return jsonify({"reply": ai_reply})

    @app.post("/api/scripts/generate")
    @auth_required(app)
    def generate_script():
        body = request.get_json(silent=True) or {}
        idea = (body.get("idea") or "").strip()
        mode = body.get("mode") or "default"
        submode = body.get("submode") or "universal"
        niche = body.get("niche") or ""
        audience = body.get("audience") or ""
        tone = body.get("tone") or ""
        goal = body.get("goal") or ""
        language = body.get("language") or "en"
        platform = body.get("platform") or "generic"
        if not idea:
            return jsonify({"error": "Idea is required"}), 400

        script = generate_script_text(
            idea=idea,
            mode=mode,
            submode=submode,
            niche=niche,
            audience=audience,
            tone=tone,
            goal=goal,
            language=language,
            platform=platform,
        )
        return jsonify({"script": script})

    @app.post("/api/scripts/save")
    @auth_required(app)
    def save_script():
        body = request.get_json(silent=True) or {}
        project_id = int(body.get("project_id") or 0)
        title = (body.get("title") or "Untitled Script").strip()
        content = (body.get("content") or "").strip()
        language = body.get("language") or "en"
        mode = body.get("mode") or "default"
        tags = body.get("tags") or []
        if not content:
            return jsonify({"error": "Script content is empty"}), 400
        if not ensure_project_owner(project_id, g.user["id"]):
            return jsonify({"error": "Project not found"}), 404

        db = get_db()
        now = utcnow()
        db.execute(
            "INSERT INTO scripts(project_id,title,content,language,mode,tags,created_at,updated_at,version) VALUES (?,?,?,?,?,?,?,?,?)",
            (project_id, title, content, language, mode, json.dumps(tags, ensure_ascii=False), now, now, 1),
        )
        script_id = db.execute("SELECT last_insert_rowid() as id").fetchone()["id"]
        db.execute(
            "INSERT INTO script_versions(script_id,version,content,created_at) VALUES (?,?,?,?)",
            (script_id, 1, content, now),
        )
        db.execute("UPDATE projects SET updated_at = ? WHERE id = ?", (now, project_id))
        db.commit()
        return jsonify({"ok": True, "id": script_id})

    @app.get("/api/projects/<int:project_id>/scripts")
    @auth_required(app)
    def list_scripts(project_id: int):
        if not ensure_project_owner(project_id, g.user["id"]):
            return jsonify({"error": "Project not found"}), 404
        rows = get_db().execute(
            "SELECT id,title,language,mode,tags,version,created_at,updated_at FROM scripts WHERE project_id = ? ORDER BY updated_at DESC",
            (project_id,),
        ).fetchall()
        data = []
        for r in rows:
            item = dict(r)
            item["tags"] = json.loads(item.get("tags") or "[]")
            data.append(item)
        return jsonify(data)

    @app.get("/api/scripts/<int:script_id>")
    @auth_required(app)
    def get_script(script_id: int):
        row = get_db().execute(
            "SELECT s.* FROM scripts s JOIN projects p ON p.id = s.project_id WHERE s.id = ? AND p.user_id = ?",
            (script_id, g.user["id"]),
        ).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404
        item = dict(row)
        item["tags"] = json.loads(item.get("tags") or "[]")
        return jsonify(item)

    @app.post("/api/scripts/<int:script_id>/version")
    @auth_required(app)
    def save_script_version(script_id: int):
        body = request.get_json(silent=True) or {}
        new_content = (body.get("content") or "").strip()
        if not new_content:
            return jsonify({"error": "Content is empty"}), 400

        db = get_db()
        row = db.execute(
            "SELECT s.id,s.project_id,s.version FROM scripts s JOIN projects p ON p.id = s.project_id WHERE s.id = ? AND p.user_id = ?",
            (script_id, g.user["id"]),
        ).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404

        new_version = int(row["version"]) + 1
        now = utcnow()
        db.execute("UPDATE scripts SET content = ?, version = ?, updated_at = ? WHERE id = ?", (new_content, new_version, now, script_id))
        db.execute(
            "INSERT INTO script_versions(script_id,version,content,created_at) VALUES (?,?,?,?)",
            (script_id, new_version, new_content, now),
        )
        db.execute("UPDATE projects SET updated_at = ? WHERE id = ?", (now, row["project_id"]))
        db.commit()
        return jsonify({"ok": True, "version": new_version})

    @app.get("/api/scripts/<int:script_id>/versions")
    @auth_required(app)
    def list_versions(script_id: int):
        rows = get_db().execute(
            "SELECT v.version,v.content,v.created_at FROM script_versions v JOIN scripts s ON s.id = v.script_id JOIN projects p ON p.id = s.project_id WHERE s.id = ? AND p.user_id = ? ORDER BY v.version DESC",
            (script_id, g.user["id"]),
        ).fetchall()
        return jsonify([dict(r) for r in rows])

    @app.delete("/api/scripts/<int:script_id>")
    @auth_required(app)
    def delete_script(script_id: int):
        db = get_db()
        row = db.execute(
            "SELECT s.project_id FROM scripts s JOIN projects p ON p.id = s.project_id WHERE s.id = ? AND p.user_id = ?",
            (script_id, g.user["id"]),
        ).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404
        db.execute("DELETE FROM script_versions WHERE script_id = ?", (script_id,))
        db.execute("DELETE FROM scripts WHERE id = ?", (script_id,))
        db.execute("UPDATE projects SET updated_at = ? WHERE id = ?", (utcnow(), row["project_id"]))
        db.commit()
        return jsonify({"ok": True})

    @app.get("/api/scripts/<int:script_id>/export")
    @auth_required(app)
    def export_script(script_id: int):
        export_format = (request.args.get("format") or "txt").lower()
        row = get_db().execute(
            "SELECT s.title,s.content,s.language,s.mode,s.updated_at FROM scripts s JOIN projects p ON p.id = s.project_id WHERE s.id = ? AND p.user_id = ?",
            (script_id, g.user["id"]),
        ).fetchone()
        if not row:
            return jsonify({"error": "Not found"}), 404

        title = row["title"]
        content = row["content"]
        meta = f"Title: {title}\nLanguage: {row['language']}\nMode: {row['mode']}\nUpdated: {row['updated_at']}\n\n"

        if export_format == "txt":
            payload = (meta + content).encode("utf-8")
            return send_file(io.BytesIO(payload), as_attachment=True, download_name=f"{safe_name(title)}.txt", mimetype="text/plain")

        if export_format == "md":
            md = f"# {title}\n\n{meta}\n{content}\n"
            return send_file(io.BytesIO(md.encode("utf-8")), as_attachment=True, download_name=f"{safe_name(title)}.md", mimetype="text/markdown")

        if export_format == "docx":
            payload = build_docx(title=title, meta=meta, content=content)
            return send_file(io.BytesIO(payload), as_attachment=True, download_name=f"{safe_name(title)}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if export_format == "pdf":
            payload = build_pdf(title=title, meta=meta, content=content)
            return send_file(io.BytesIO(payload), as_attachment=True, download_name=f"{safe_name(title)}.pdf", mimetype="application/pdf")

        return jsonify({"error": "Unsupported format"}), 400

    @app.post("/api/ideas/suggest")
    @auth_required(app)
    def suggest_ideas():
        body = request.get_json(silent=True) or {}
        niche = (body.get("niche") or "general").strip()
        audience = (body.get("audience") or "broad audience").strip()
        language = (body.get("language") or "en").lower()
        direction = (body.get("direction") or "default").strip().lower()
        ideas = suggest_idea_list(niche, audience, language, direction)
        return jsonify({"ideas": ideas})

    @app.post("/api/qa/analyze")
    @auth_required(app)
    def qa_analyze():
        body = request.get_json(silent=True) or {}
        analysis_type = body.get("analysis_type") or "platform_check"
        script = body.get("script") or ""
        if not script.strip():
            return jsonify({"error": "Script is empty"}), 400

        if analysis_type == "budget_estimate":
            budget = body.get("budget") or "mid"
            region = body.get("region") or "US"
            breakdown = estimate_budget(script, budget, region)
            return jsonify(breakdown)

        platform = (body.get("platform") or "Netflix").strip()
        result = check_platform_compliance(script, platform)
        return jsonify(result)

    return app


class SimpleRateLimiter:
    def __init__(self, per_minute: int):
        self.per_minute = per_minute
        self.requests = defaultdict(deque)

    def allow(self, key: str) -> bool:
        now = dt.datetime.now(dt.UTC).timestamp()
        q = self.requests[key]
        while q and now - q[0] > 60:
            q.popleft()
        if len(q) >= self.per_minute:
            return False
        q.append(now)
        return True


def get_db() -> sqlite3.Connection:
    if "db" not in g:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        g.db = conn
    return g.db


def init_db() -> None:
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA foreign_keys = ON")
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            name TEXT NOT NULL,
            password_hash TEXT NOT NULL,
            is_admin INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            status TEXT NOT NULL,
            tags TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS scripts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            content TEXT NOT NULL,
            language TEXT NOT NULL,
            mode TEXT NOT NULL,
            tags TEXT,
            version INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS script_versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            script_id INTEGER NOT NULL,
            version INTEGER NOT NULL,
            content TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(script_id) REFERENCES scripts(id) ON DELETE CASCADE
        );
        """
    )
    ensure_column(conn, "users", "is_admin", "is_admin INTEGER NOT NULL DEFAULT 0")
    conn.commit()
    conn.close()


def ensure_column(conn: sqlite3.Connection, table: str, column: str, ddl_fragment: str) -> None:
    cols = [row[1] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()]
    if column not in cols:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {ddl_fragment}")


def is_admin_email(email: str) -> bool:
    raw = os.getenv("ADMIN_EMAILS", "")
    admins = {item.strip().lower() for item in raw.split(",") if item.strip()}
    return email.strip().lower() in admins


def auth_required(app: Flask):
    def decorator(fn):
        @wraps(fn)
        def wrapper(*args, **kwargs):
            token = None
            auth = request.headers.get("Authorization", "")
            if auth.startswith("Bearer "):
                token = auth.split(" ", 1)[1].strip()
            if not token:
                return jsonify({"error": "Unauthorized"}), 401
            try:
                user_data = read_token(app, token)
            except SignatureExpired:
                return jsonify({"error": "Token expired"}), 401
            except BadSignature:
                return jsonify({"error": "Invalid token"}), 401
            db = get_db()
            row = db.execute(
                "SELECT id,email,name,is_admin FROM users WHERE id = ?",
                (user_data["uid"],),
            ).fetchone()
            if not row:
                return jsonify({"error": "Unauthorized"}), 401
            g.user = {
                "id": row["id"],
                "email": row["email"],
                "name": row["name"],
                "is_admin": int(row["is_admin"] or 0),
            }
            return fn(*args, **kwargs)

        return wrapper

    return decorator


def sign_token(app: Flask, data: dict) -> str:
    s = URLSafeTimedSerializer(app.config["SECRET_KEY"], salt="leadscript-auth")
    return s.dumps(data)


def read_token(app: Flask, token: str) -> dict:
    s = URLSafeTimedSerializer(app.config["SECRET_KEY"], salt="leadscript-auth")
    return s.loads(token, max_age=app.config["TOKEN_MAX_AGE"])


def ensure_project_owner(project_id: int, user_id: int) -> bool:
    row = get_db().execute("SELECT id FROM projects WHERE id = ? AND user_id = ?", (project_id, user_id)).fetchone()
    return bool(row)


def row_to_project(row: sqlite3.Row) -> dict:
    item = dict(row)
    item["tags"] = json.loads(item.get("tags") or "[]")
    return item


def utcnow() -> str:
    return dt.datetime.now(dt.UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def safe_name(name: str) -> str:
    cleaned = "".join(ch for ch in name if ch.isalnum() or ch in ("-", "_", " ")).strip()
    return cleaned.replace(" ", "_") or "script"


def build_docx(title: str, meta: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in meta.strip().splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph("")
    for line in content.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(title: str, meta: str, content: str) -> bytes:
    buf = io.BytesIO()
    pdf = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x = 48
    y = height - 48

    def write_line(text: str):
        nonlocal y
        pdf.drawString(x, y, text[:120])
        y -= 16
        if y < 48:
            pdf.showPage()
            y = height - 48

    pdf.setFont("Helvetica-Bold", 15)
    write_line(title)
    pdf.setFont("Helvetica", 10)
    for line in meta.strip().splitlines():
        write_line(line)
    write_line("")
    for line in content.splitlines():
        write_line(line)
    pdf.save()
    return buf.getvalue()


def hash_short(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:8]


def generate_ai_reply(
    text: str,
    style: str,
    detail: str,
    focus: str,
    creativity: str,
    model: str = "gpt-5.3-codex",
    temperature: float = 0.7,
    max_tokens: int = 350,
    response_mode: str = "normal",
    system_prompt: str = "",
    history: list[dict] | None = None,
) -> str:
    # ChatGPT-like behavior for simple greetings / small talk.
    small_talk = detect_small_talk(text)
    if small_talk:
        return small_talk

    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    selected_model = model or os.getenv("OPENAI_MODEL", "gpt-5.3-codex")
    history = history or []
    mode_hints = {
        "concise": "Keep response concise and compact.",
        "normal": "Balance concise advice with enough practical detail.",
        "deep": "Provide deeper strategic explanation with examples.",
    }
    model_whitelist = {
        "gpt-5.3-codex",
        "gpt-4.1-mini",
        "gpt-4.1",
        "gpt-4o-mini",
        "gpt-4o",
    }
    if selected_model not in model_whitelist:
        selected_model = os.getenv("OPENAI_MODEL", "gpt-5.3-codex")

    if api_key and OpenAI is not None:
        try:
            client = OpenAI(api_key=api_key)
            history_text = "\n".join(
                f"{h.get('role','user')}: {h.get('content','')}" for h in history[-16:]
            )
            sys_text = system_prompt if system_prompt else "You are AliceAI for creators."
            prompt = (
                f"{sys_text}\n"
                "You are an assistant for scriptwriters and content creators. "
                "Provide practical guidance with actionable bullets.\n"
                f"Mode hint: {mode_hints.get(response_mode, mode_hints['normal'])}\n"
                f"Settings: style={style}, detail={detail}, focus={focus}, creativity={creativity}.\n"
                f"Recent context:\n{history_text}\n"
                f"User request: {text}\n"
            )
            resp = client.responses.create(
                model=selected_model,
                input=prompt,
                temperature=temperature,
                max_output_tokens=max_tokens,
            )
            if hasattr(resp, "output_text") and resp.output_text:
                return resp.output_text.strip()
        except Exception:
            pass

    # Deterministic fallback (works without external API).
    tips = [
        "Define one clear outcome for this script.",
        "Start with a hook in the first 2 seconds.",
        "Use one core conflict and one turning point.",
        "End with one concrete CTA.",
    ]
    return (
        f"AliceAI quick guidance for: \"{text}\"\n"
        f"Settings: style={style}, detail={detail}, focus={focus}, creativity={creativity}, "
        f"model={selected_model}, temperature={temperature}, max_tokens={max_tokens}, mode={response_mode}.\n\n"
        + "\n".join(f"{i+1}. {t}" for i, t in enumerate(tips))
        + f"\n\nRef: {hash_short(text + style + detail + focus + creativity)}"
    )


def detect_small_talk(text: str) -> str | None:
    raw = (text or "").strip()
    if not raw:
        return None
    low = raw.lower()
    simple = {
        "hi",
        "hello",
        "hey",
        "yo",
        "hola",
        "sup",
        "start",
        "ping",
        "привет",
        "здарова",
        "здравствуйте",
        "ку",
        "алло",
        "хай",
    }
    how_are_you = {
        "how are you",
        "как дела",
        "как ты",
        "че как",
    }

    if low in simple or len(low) <= 5 and low.isalpha():
        return "Hi! What would you like to work on today?"
    if any(p in low for p in how_are_you):
        return "I am doing great. Want to draft a script, improve one, or generate fresh ideas?"
    if low in {"thanks", "thank you", "спасибо", "благодарю"}:
        return "You are welcome. If you want, send a topic and I will turn it into a strong script."
    return None


def clamp_float(value, low: float, high: float, default: float) -> float:
    try:
        num = float(value)
    except Exception:
        return default
    return max(low, min(high, num))


def clamp_int(value, low: int, high: int, default: int) -> int:
    try:
        num = int(value)
    except Exception:
        return default
    return max(low, min(high, num))


def generate_script_text(**kwargs) -> str:
    idea = kwargs["idea"]
    mode = kwargs["mode"]
    submode = kwargs["submode"]
    niche = kwargs["niche"]
    audience = kwargs["audience"]
    tone = kwargs["tone"]
    goal = kwargs["goal"]
    language = kwargs["language"]
    platform = kwargs["platform"]
    mode_key = (mode or "default").lower()

    if language == "ru":
        mode_blocks_ru = {
            "business": (
                "1) Проблема бизнеса:\n"
                "- Какая точка потерь денег/времени есть сейчас.\n\n"
                "2) Решение и монетизация:\n"
                "- Как продукт решает боль и за счет чего окупается.\n\n"
                "3) MVP-план (14 дней):\n"
                "- Гипотеза -> тест -> метрика успеха.\n\n"
                "4) CTA для клиента/инвестора:\n"
                "- Следующий шаг: демо, звонок или пилот.\n"
            ),
            "gamedev": (
                "1) Core Loop:\n"
                "- Действие игрока -> награда -> прогресс.\n\n"
                "2) USP игры:\n"
                "- Что отличает игру от конкурентов.\n\n"
                "3) Контент-план релиза:\n"
                "- Вертикальный срез, плейтест, soft launch.\n\n"
                "4) Метрики:\n"
                "- D1/D7 retention, CPI, conversion в payers.\n"
            ),
            "programming": (
                "1) Техническая задача:\n"
                "- Что реализуем и какие ограничения.\n\n"
                "2) Архитектура решения:\n"
                "- Модули, API-контракты, зависимости.\n\n"
                "3) План разработки:\n"
                "- Шаги имплементации + тесты + edge-cases.\n\n"
                "4) Результат и деплой:\n"
                "- Проверка качества и выпуск в production.\n"
            ),
        }
        body_ru = mode_blocks_ru.get(
            mode_key,
            "1) Хук (0-2 сек):\n"
            f"- {idea}: контрастный факт/боль аудитории.\n\n"
            "2) Основной блок (3-20 сек):\n"
            "- Проблема -> решение -> мини-доказательство (кейс/цифра).\n\n"
            "3) Финал (20-30 сек):\n"
            "- Четкий CTA + дедлайн/триггер.\n",
        )
        return (
            f"СЦЕНАРИЙ ({mode}/{submode})\n"
            f"Идея: {idea}\nНиша: {niche}\nАудитория: {audience}\nТон: {tone}\nЦель: {goal}\nПлатформа: {platform}\n\n"
            + body_ru
        )

    mode_blocks_en = {
        "business": (
            "1) Business pain statement:\n"
            "- Name the revenue/time bottleneck clearly.\n\n"
            "2) Solution and monetization:\n"
            "- Explain value creation and pricing logic.\n\n"
            "3) 14-day MVP sprint:\n"
            "- Hypothesis -> experiment -> success metric.\n\n"
            "4) CTA for buyer/investor:\n"
            "- Book demo, pilot, or strategic call.\n"
        ),
        "gamedev": (
            "1) Core loop:\n"
            "- Player action -> reward -> progression.\n\n"
            "2) Game USP:\n"
            "- Why this game stands out in the genre.\n\n"
            "3) Release content plan:\n"
            "- Vertical slice, playtest, soft launch.\n\n"
            "4) Performance metrics:\n"
            "- D1/D7 retention, CPI, payer conversion.\n"
        ),
        "programming": (
            "1) Technical objective:\n"
            "- Define feature scope and constraints.\n\n"
            "2) Architecture outline:\n"
            "- Modules, API contracts, dependencies.\n\n"
            "3) Implementation plan:\n"
            "- Steps, tests, and edge-case handling.\n\n"
            "4) Delivery and rollout:\n"
            "- QA checklist and production release path.\n"
        ),
    }
    body_en = mode_blocks_en.get(
        mode_key,
        "1) Hook (0-2s):\n"
        f"- Lead with tension tied to: {idea}.\n\n"
        "2) Body (3-20s):\n"
        "- Problem -> solution -> proof (micro case or number).\n\n"
        "3) Ending (20-30s):\n"
        "- One CTA + urgency trigger.\n",
    )
    return (
        f"SCRIPT ({mode}/{submode})\n"
        f"Idea: {idea}\nNiche: {niche}\nAudience: {audience}\nTone: {tone}\nGoal: {goal}\nPlatform: {platform}\n\n"
        + body_en
    )


def suggest_idea_list(niche: str, audience: str, language: str, direction: str = "default") -> list[str]:
    direction_hint_ru = {
        "business": "роста бизнеса",
        "gamedev": "разработки игры",
        "programming": "разработки ПО",
    }.get(direction, "контента")
    direction_hint_en = {
        "business": "business growth",
        "gamedev": "game development",
        "programming": "software engineering",
    }.get(direction, "content strategy")
    if language == "ru":
        return [
            f"Разбор 3 ошибок {direction_hint_ru} в нише {niche} для {audience}",
            f"Кейс: как получить результат в {niche} за 7 дней ({direction_hint_ru})",
            f"Мифы о {niche}: что реально работает в 2026 в теме {direction_hint_ru}",
            f"Сравнение подходов: быстрый vs стратегический в {niche} ({direction_hint_ru})",
            f"Чеклист перед стартом в {niche} для задачи {direction_hint_ru}",
        ]
    return [
        f"Top 3 {direction_hint_en} mistakes in {niche} for {audience}",
        f"Case study: achieving results in {niche} within 7 days ({direction_hint_en})",
        f"Myths in {niche}: what actually works in 2026 for {direction_hint_en}",
        f"Quick wins vs long-term strategy in {niche} ({direction_hint_en})",
        f"Pre-launch checklist for {niche} with {direction_hint_en} focus",
    ]


def check_platform_compliance(script: str, platform: str) -> dict:
    platform_rules = {
        "Netflix": ["character", "conflict", "arc", "visual", "dialogue"],
        "YouTube": ["hook", "retention", "cta", "chapters", "value"],
        "TikTok": ["hook", "2 sec", "trend", "cta", "keyword"],
        "Kion": ["character", "culture", "arc", "dialogue"],
        "Start": ["hook", "character", "conflict", "payoff"],
    }
    rules = platform_rules.get(platform, platform_rules["Netflix"])
    text = script.lower()
    found = [r for r in rules if r.lower() in text]
    missing = [r for r in rules if r.lower() not in text]
    score = int((len(found) / len(rules)) * 100)
    verdict = "pass" if score >= 70 else "revise"
    return {
        "platform": platform,
        "score": score,
        "verdict": verdict,
        "matched": found,
        "missing": missing,
        "recommendations": [
            "Strengthen hook in opening.",
            "Ensure character motivation is explicit.",
            "Add one measurable CTA.",
        ],
    }


def estimate_budget(script: str, budget_level: str, region: str) -> dict:
    scene_markers = ["scene", "сцена", "int.", "ext."]
    lines = [ln.strip() for ln in script.splitlines() if ln.strip()]
    scenes = []
    cur = []
    for ln in lines:
        low = ln.lower()
        if any(m in low for m in scene_markers) and cur:
            scenes.append(cur)
            cur = [ln]
        else:
            cur.append(ln)
    if cur:
        scenes.append(cur)
    if not scenes:
        scenes = [lines or ["Scene draft"]]

    base_map = {"low": 250, "mid": 700, "high": 1800}
    region_mult = {"US": 1.0, "EU": 0.9, "CIS": 0.65, "LATAM": 0.7}
    base = base_map.get(budget_level, 700)
    rm = region_mult.get(region.upper(), 1.0)

    detail = []
    total = 0.0
    for idx, sc in enumerate(scenes, start=1):
        text = " ".join(sc).lower()
        complexity = 1.0
        if any(k in text for k in ["crowd", "chase", "stunt", "rain", "night", "drone"]):
            complexity += 0.45
        if any(k in text for k in ["vfx", "cgi", "3d"]):
            complexity += 0.6
        if any(k in text for k in ["dialogue", "talk"]):
            complexity += 0.1

        scene_cost = round(base * complexity * rm, 2)
        lighting = round(scene_cost * 0.18, 2)
        cast = round(scene_cost * 0.32, 2)
        equipment = round(scene_cost * 0.27, 2)
        post = round(scene_cost * 0.23, 2)
        total += scene_cost
        detail.append(
            {
                "scene": idx,
                "cost": scene_cost,
                "complexity": round(complexity, 2),
                "items": {
                    "lighting": lighting,
                    "cast": cast,
                    "equipment": equipment,
                    "post": post,
                },
            }
        )

    contingency = round(total * 0.12, 2)
    grand_total = round(total + contingency, 2)
    return {
        "budget_level": budget_level,
        "region": region.upper(),
        "scenes": len(detail),
        "breakdown": detail,
        "subtotal": round(total, 2),
        "contingency": contingency,
        "total": grand_total,
    }


app = create_app()


@app.teardown_appcontext
def close_db(_exc):
    db = g.pop("db", None)
    if db is not None:
        db.close()


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5050"))
    app.run(host="127.0.0.1", port=port, debug=False)
